#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Генератор Word-документов из Excel-файлов
Скрипт для создания таблиц дефектов в формате .docx из данных Excel

Автор: Автоматически сгенерированный скрипт
Поддерживаемые ОС: Windows, macOS
"""

import os
import sys
import subprocess
import re
from datetime import datetime
from pathlib import Path
from typing import List, Tuple, Optional

try:
    from openpyxl import load_workbook
    from openpyxl.utils import get_column_letter
except ImportError:
    print("ОШИБКА: Библиотека openpyxl не установлена.")
    print("Установите её командой: pip install openpyxl")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
except ImportError:
    print("ОШИБКА: Библиотека python-docx не установлена.")
    print("Установите её командой: pip install python-docx")
    sys.exit(1)


# ============================================================================
# КОНФИГУРАЦИЯ ПУТЕЙ И ПАРАМЕТРОВ
# ============================================================================

# Базовый путь к скрипту
BASE_DIR = Path(__file__).parent.absolute()

# Пути к директориям
DATA_DIR = BASE_DIR / "data"
OUTPUT_DIR = BASE_DIR / "output"
TEMPLATES_DIR = BASE_DIR / "templates"
FONTS_DIR = BASE_DIR / "fonts"

# Заголовки таблицы (6 столбцов: № п/п + 5 основных столбцов)
TABLE_HEADERS = [
    "№ п/п",
    "Наименование дефекта/повреждения",
    "Место расположения",
    "Номер фото из фотоматериалов",
    "Контролируемый параметр дефекта (глубина, ширина раскрытия, твердость, прочность, сплошность)",
    "Примечание"
]

# Номера столбцов для второй строки заголовков (все столбцы пронумерованы)
COLUMN_NUMBERS = ["1", "2", "3", "4", "5", "6"]

# Параметры форматирования
FONT_NAME = "Roboto"  # Используется Roboto, если доступен, иначе Liberation Sans
FONT_SIZE = Pt(11)
TABLE_BORDER_WIDTH = Pt(0.5)
TABLE_BORDER_COLOR = RGBColor(0, 0, 0)  # Чёрный


# ============================================================================
# ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ
# ============================================================================

def create_directories():
    """Создаёт необходимые директории, если они не существуют."""
    directories = [DATA_DIR, OUTPUT_DIR, TEMPLATES_DIR, FONTS_DIR]
    for directory in directories:
        directory.mkdir(exist_ok=True)
        print(f"✓ Проверена директория: {directory}")


def get_excel_files() -> List[Path]:
    """
    Сканирует папку /data и возвращает список доступных Excel-файлов.
    
    Returns:
        List[Path]: Список путей к .xlsx и .xls файлам
    """
    excel_files = []
    
    if not DATA_DIR.exists():
        print(f"ОШИБКА: Папка {DATA_DIR} не существует. Создаю...")
        create_directories()
        return excel_files
    
    # Ищем файлы с расширениями .xlsx и .xls
    extensions = ['.xlsx', '.xls']
    for ext in extensions:
        excel_files.extend(DATA_DIR.glob(f'*{ext}'))
    
    return sorted(excel_files)


def select_excel_file(excel_files: List[Path]) -> Optional[Path]:
    """
    Позволяет пользователю выбрать Excel-файл из списка.
    
    Args:
        excel_files: Список путей к Excel-файлам
        
    Returns:
        Path: Выбранный файл или None, если выбор отменён
    """
    if not excel_files:
        print("\n" + "="*60)
        print("⚠ ВНИМАНИЕ: Excel-файлы не найдены")
        print("="*60)
        print(f"\n📁 Папка для Excel-файлов: {DATA_DIR}")
        print("\n📋 Что нужно сделать:")
        print("   1. Создайте Excel-файл (.xlsx или .xls) или используйте существующий")
        print("   2. Поместите его в папку 'data' (скопируйте файл в эту папку)")
        print("   3. Запустите скрипт снова")
        print("\n💡 Рекомендуемая структура Excel-файла:")
        print("   - Наименование дефекта/повреждения")
        print("   - Место расположения")
        print("   - Номер фото из фотоматериалов")
        print("   - Контролируемый параметр дефекта")
        print("   - Примечание")
        print("\n" + "="*60)
        return None
    
    print("\n" + "="*60)
    print("ДОСТУПНЫЕ EXCEL-ФАЙЛЫ:")
    print("="*60)
    
    for idx, file_path in enumerate(excel_files, start=1):
        print(f"{idx}. {file_path.name}")
    
    while True:
        try:
            choice = input(f"\nВыберите номер файла (1-{len(excel_files)}): ").strip()
            file_idx = int(choice) - 1
            
            if 0 <= file_idx < len(excel_files):
                selected_file = excel_files[file_idx]
                print(f"✓ Выбран файл: {selected_file.name}")
                return selected_file
            else:
                print(f"❌ Неверный номер. Введите число от 1 до {len(excel_files)}")
        except ValueError:
            print("❌ Введите корректное число")
        except KeyboardInterrupt:
            print("\n\nОперация отменена пользователем.")
            return None


def find_column_index(headers: List[str], search_names: List[str]) -> Optional[int]:
    """
    Находит индекс столбца по различным вариантам названий.
    
    Args:
        headers: Список заголовков столбцов из Excel
        search_names: Список вариантов названий для поиска
        
    Returns:
        Optional[int]: Индекс найденного столбца или None
    """
    # Нормализуем строки для поиска (убираем лишние пробелы, приводим к нижнему регистру)
    def normalize(text):
        """Нормализует текст: убирает лишние пробелы, приводит к нижнему регистру"""
        text = str(text).strip()
        # Заменяем множественные пробелы на один
        text = re.sub(r'\s+', ' ', text)
        return text.lower()
    
    normalized_headers = [normalize(h) for h in headers]
    
    for search_name in search_names:
        normalized_search = normalize(search_name)
        
        # Точное совпадение
        if normalized_search in normalized_headers:
            return normalized_headers.index(normalized_search)
        
        # Поиск по подстроке (содержит) - более гибкий поиск
        for idx, header in enumerate(normalized_headers):
            # Удаляем все пробелы для более гибкого сравнения
            if (normalized_search.replace(' ', '') in header.replace(' ', '') or 
                header.replace(' ', '') in normalized_search.replace(' ', '')):
                return idx
    
    return None


def read_excel_file(file_path: Path, sheet_name: Optional[str] = None) -> Tuple[List[List[str]], Optional[str], dict, Optional[int]]:
    """
    Читает данные из Excel-файла и определяет маппинг столбцов.
    
    Args:
        file_path: Путь к Excel-файлу
        sheet_name: Имя листа для чтения (если None, используется первый лист)
        
    Returns:
        Tuple[List[List[str]], Optional[str], dict, Optional[int]]: 
        (Данные, имя листа, маппинг столбцов, индекс столбца конструкций)
        Маппинг: {word_column_index: excel_column_index} или {word_column_index: None}
    """
    try:
        workbook = load_workbook(file_path, data_only=True)
        
        # Если лист не указан, используем первый
        if sheet_name is None:
            sheet_name = workbook.sheetnames[0]
        
        sheet = workbook[sheet_name]
        
        # Если есть несколько листов, предлагаем выбрать
        if len(workbook.sheetnames) > 1:
            print(f"\nВ файле найдено {len(workbook.sheetnames)} лист(ов):")
            for idx, name in enumerate(workbook.sheetnames, start=1):
                marker = " ← выбран" if name == sheet_name else ""
                print(f"  {idx}. {name}{marker}")
            
            choice = input(f"Использовать лист '{sheet_name}'? (Enter = да, или введите номер другого листа): ").strip()
            if choice:
                try:
                    sheet_idx = int(choice) - 1
                    if 0 <= sheet_idx < len(workbook.sheetnames):
                        sheet_name = workbook.sheetnames[sheet_idx]
                        sheet = workbook[sheet_name]
                    else:
                        print(f"⚠ Неверный номер. Используется лист '{sheet_name}'")
                except ValueError:
                    print(f"⚠ Используется лист '{sheet_name}'")
        
        # Читаем первую строку как заголовки
        first_row = next(sheet.iter_rows(values_only=True))
        headers = [str(cell).strip() if cell is not None else "" for cell in first_row]
        
        print(f"\n📋 Найденные заголовки столбцов в Excel:")
        for idx, header in enumerate(headers):
            if header:
                print(f"   Столбец {idx + 1}: {header}")
        
        # Определяем маппинг столбцов Excel к столбцам Word-таблицы
        # Индексы Word-таблицы: 0=№ п/п, 1=Наименование, 2=Место, 3=Фото, 4=Параметр, 5=Примечание
        # В маппинге: ключ - это внутренний индекс (1-5), значение - индекс столбца в Excel
        column_mapping = {}
        
        # Столбец 1: Наименование дефекта/повреждения (индекс 1 в Word-таблице)
        excel_col = find_column_index(headers, [
            "Наименование дефекта / повреждения",
            "Наименование дефекта/повреждения",
            "Наименование дефекта",
            "Дефект",
            "Повреждение"
        ])
        column_mapping[1] = excel_col
        if excel_col is not None:
            print(f"✓ Найден столбец 'Наименование дефекта/повреждения': столбец {excel_col + 1} ({headers[excel_col]})")
        else:
            print(f"⚠ Столбец 'Наименование дефекта/повреждения' не найден, будет использован первый столбец")
        
        # Столбец 2: Место расположения - оставляем пустым (не заполняем из Excel)
        column_mapping[2] = None
        
        # Находим столбец с наименованием конструкций (для группировки)
        construction_col = find_column_index(headers, [
            "Наименование конструкций",
            "Конструкции",
            "Конструкция",
            "Тип конструкции",
            "Элемент конструкции"
        ])
        if construction_col is not None:
            print(f"✓ Найден столбец 'Наименование конструкций': столбец {construction_col + 1} ({headers[construction_col]})")
        
        # Столбец 3: Номер фото из фотоматериалов
        excel_col = find_column_index(headers, [
            "Номер фото из фотоматериалов",
            "Номер фото",
            "Фото",
            "Фотоматериалы",
            "Номер фото из фото"
        ])
        column_mapping[3] = excel_col
        
        # Столбец 4: Контролируемый параметр дефекта
        excel_col = find_column_index(headers, [
            "Контролируемый параметр дефекта",
            "Параметр дефекта",
            "Контролируемый параметр",
            "Параметр",
            "Параметры"
        ])
        column_mapping[4] = excel_col
        
        # Столбец 5: Примечание
        excel_col = find_column_index(headers, [
            "Примечание",
            "Примечания",
            "Комментарий",
            "Замечания"
        ])
        column_mapping[5] = excel_col
        
        # Читаем все данные из листа (начиная со второй строки, так как первая - заголовки)
        data = []
        for row_idx, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            if row_idx == 1:
                continue  # Пропускаем заголовки
            # Преобразуем значения в строки, заменяя None на пустую строку
            row_data = [str(cell) if cell is not None else "" for cell in row]
            data.append(row_data)
        
        workbook.close()
        print(f"✓ Загружено {len(data)} строк данных из листа '{sheet_name}'")
        return data, sheet_name, column_mapping, construction_col
        
    except Exception as e:
        print(f"❌ Ошибка при чтении Excel-файла: {e}")
        import traceback
        traceback.print_exc()
        return [], None, {}, None


def display_rows_preview(data: List[List[str]], max_rows: int = 20) -> None:
    """
    Отображает превью строк для пользователя.
    
    Args:
        data: Данные из Excel
        max_rows: Максимальное количество строк для отображения
    """
    if not data:
        print("⚠ Нет данных для отображения")
        return
    
    print("\n" + "="*60)
    print("ПРЕВЬЮ ДАННЫХ (первые столбцы):")
    print("="*60)
    
    # Показываем первые столбцы каждой строки
    for idx, row in enumerate(data[:max_rows], start=1):
        preview = " | ".join(str(cell)[:30] for cell in row[:2])  # Первые 2 столбца
        print(f"{idx:3d}. {preview}")
    
    if len(data) > max_rows:
        print(f"... и ещё {len(data) - max_rows} строк(и)")
    
    print("="*60)


def parse_row_selection(selection: str, total_rows: int) -> List[int]:
    """
    Парсит строку выбора пользователя и возвращает список индексов строк.
    
    Поддерживаемые форматы:
    - "1" - одна строка
    - "2-5" - диапазон
    - "1,3,7" - несколько через запятую
    - "all" - все строки
    
    Args:
        selection: Строка выбора пользователя
        total_rows: Общее количество доступных строк
        
    Returns:
        List[int]: Список индексов строк (0-based)
    """
    selection = selection.strip().lower()
    
    if selection == "all":
        return list(range(total_rows))
    
    selected_indices = []
    
    # Разделяем по запятой
    parts = selection.split(',')
    
    for part in parts:
        part = part.strip()
        
        if '-' in part:
            # Диапазон (например, "2-5")
            try:
                start, end = part.split('-', 1)
                start_idx = int(start.strip()) - 1  # Преобразуем в 0-based
                end_idx = int(end.strip())  # Пользователь вводит 1-based
                
                if start_idx < 0:
                    start_idx = 0
                if end_idx > total_rows:
                    end_idx = total_rows
                
                selected_indices.extend(range(start_idx, end_idx))
            except ValueError:
                print(f"⚠ Некорректный диапазон: {part}")
        else:
            # Одна строка
            try:
                row_idx = int(part) - 1  # Преобразуем в 0-based
                if 0 <= row_idx < total_rows:
                    selected_indices.append(row_idx)
                else:
                    print(f"⚠ Строка {part} вне диапазона (1-{total_rows})")
            except ValueError:
                print(f"⚠ Некорректный номер строки: {part}")
    
    # Удаляем дубликаты и сортируем
    selected_indices = sorted(set(selected_indices))
    return selected_indices


def get_available_constructions(data_rows: List[List[str]], construction_col: Optional[int]) -> List[str]:
    """
    Получает список доступных конструкций из данных Excel.
    
    Args:
        data_rows: Все строки данных из Excel
        construction_col: Индекс столбца с наименованием конструкций или None
        
    Returns:
        List[str]: Список уникальных конструкций (отсортированный)
    """
    if construction_col is None:
        return []
    
    constructions = set()
    for row in data_rows:
        if construction_col < len(row):
            construction_name = str(row[construction_col]).strip()
            if construction_name:
                constructions.add(construction_name)
    
    return sorted(list(constructions))


def select_constructions(data_rows: List[List[str]], construction_col: Optional[int]) -> Optional[List[str]]:
    """
    Позволяет пользователю выбрать конструкции, где описываются дефекты.
    
    Args:
        data_rows: Все строки данных из Excel
        construction_col: Индекс столбца с наименованием конструкций или None
        
    Returns:
        Optional[List[str]]: Список выбранных конструкций или None (если столбец не найден)
    """
    if construction_col is None:
        print("\n" + "="*60)
        print("⚠ ВНИМАНИЕ: Столбец 'Наименование конструкций' не найден в Excel")
        print("="*60)
        print("В файле отсутствует столбец с наименованием конструкций.")
        print("Все строки будут обработаны без фильтрации по конструкциям.")
        print("="*60)
        return None
    
    constructions = get_available_constructions(data_rows, construction_col)
    
    if not constructions:
        print("\n" + "="*60)
        print("⚠ ВНИМАНИЕ: В данных не найдено конструкций")
        print("="*60)
        print("В выбранных строках нет данных о конструкциях.")
        print("Все строки будут обработаны.")
        print("="*60)
        return None
    
    print("\n" + "="*60)
    print("ВЫБОР КОНСТРУКЦИЙ ДЛЯ ОПИСАНИЯ ДЕФЕКТОВ")
    print("="*60)
    print("Выберите конструкции, для которых нужно создать таблицу дефектов:")
    print("="*60)
    
    for idx, construction in enumerate(constructions, start=1):
        print(f"  {idx}. {construction}")
    
    print("="*60)
    print("\n💡 ИНСТРУКЦИЯ:")
    print("  • Одна конструкция: введите номер (например: 1)")
    print("  • Несколько конструкций: введите номера через запятую (например: 1,3,5)")
    print("  • Все конструкции: введите 'all'")
    print("  • Диапазон: введите диапазон (например: 1-3)")
    print("="*60)
    
    while True:
        try:
            selection = input(f"\n👉 Выберите конструкции для обработки (1-{len(constructions)}): ").strip()
            
            if not selection:
                print("❌ Пожалуйста, введите выбор конструкций")
                continue
            
            if selection.lower() == "all":
                selected_constructions = constructions
                print(f"\n✓ Выбраны ВСЕ конструкции ({len(selected_constructions)}):")
                for constr in selected_constructions:
                    print(f"   • {constr}")
                return selected_constructions
            
            selected_indices = []
            parts = selection.split(',')
            
            for part in parts:
                part = part.strip()
                
                # Обработка диапазонов (например, "1-3")
                if '-' in part:
                    try:
                        start, end = part.split('-', 1)
                        start_idx = int(start.strip()) - 1
                        end_idx = int(end.strip())
                        
                        if start_idx < 0:
                            start_idx = 0
                        if end_idx > len(constructions):
                            end_idx = len(constructions)
                        
                        selected_indices.extend(range(start_idx, end_idx))
                    except ValueError:
                        print(f"⚠ Некорректный диапазон: {part}")
                else:
                    # Одна конструкция
                    try:
                        idx = int(part) - 1
                        if 0 <= idx < len(constructions):
                            selected_indices.append(idx)
                        else:
                            print(f"⚠ Номер {part} вне диапазона (1-{len(constructions)})")
                    except ValueError:
                        print(f"⚠ Некорректный номер: {part}")
            
            if selected_indices:
                selected_constructions = [constructions[i] for i in sorted(set(selected_indices))]
                print(f"\n✓ Выбрано {len(selected_constructions)} конструкций для обработки:")
                for constr in selected_constructions:
                    print(f"   • {constr}")
                print(f"\nБудут обработаны только строки с дефектами для этих конструкций.")
                return selected_constructions
            else:
                print("❌ Не выбрано ни одной конструкции. Попробуйте снова.")
                print(f"   Введите номер от 1 до {len(constructions)} или 'all' для всех.")
                
        except KeyboardInterrupt:
            print("\n\n⚠ Операция отменена пользователем.")
            return None


def filter_rows_by_constructions(data_rows: List[List[str]], row_indices: List[int], 
                                 construction_col: Optional[int], 
                                 selected_constructions: Optional[List[str]]) -> List[int]:
    """
    Фильтрует строки по выбранным конструкциям.
    
    Args:
        data_rows: Все строки данных из Excel
        row_indices: Исходные индексы строк
        construction_col: Индекс столбца с наименованием конструкций или None
        selected_constructions: Список выбранных конструкций или None
        
    Returns:
        List[int]: Отфильтрованные индексы строк
    """
    if selected_constructions is None or construction_col is None:
        return row_indices
    
    filtered_indices = []
    for row_idx in row_indices:
        if row_idx < len(data_rows):
            row = data_rows[row_idx]
            if construction_col < len(row):
                construction_name = str(row[construction_col]).strip()
                if construction_name in selected_constructions:
                    filtered_indices.append(row_idx)
    
    return filtered_indices


def select_rows(data: List[List[str]]) -> List[int]:
    """
    Позволяет пользователю выбрать строки для обработки.
    
    Args:
        data: Данные из Excel
        
    Returns:
        List[int]: Список индексов выбранных строк (0-based)
    """
    if not data:
        return []
    
    display_rows_preview(data)
    
    print("\n" + "="*60)
    print("ВЫБОР СТРОК ДЛЯ ОБРАБОТКИ:")
    print("="*60)
    print("Примеры ввода:")
    print("  • Одна строка: 1")
    print("  • Диапазон: 2-5")
    print("  • Несколько: 1,3,7")
    print("  • Все строки: all")
    print("="*60)
    
    while True:
        try:
            selection = input(f"\nВведите номера строк (1-{len(data)}): ").strip()
            
            if not selection:
                print("❌ Введите выбор строк")
                continue
            
            selected_indices = parse_row_selection(selection, len(data))
            
            if selected_indices:
                print(f"\n✓ Выбрано {len(selected_indices)} строк(и): {selected_indices[0]+1}-{selected_indices[-1]+1}")
                return selected_indices
            else:
                print("❌ Не выбрано ни одной строки. Попробуйте снова.")
                
        except KeyboardInterrupt:
            print("\n\nОперация отменена пользователем.")
            return []


def get_font_name() -> str:
    """
    Определяет доступный шрифт (Roboto или Liberation Sans).
    
    Returns:
        str: Имя шрифта для использования
    """
    # В python-docx мы можем указать любой шрифт,
    # система попытается его использовать при наличии
    # Roboto более современный, поэтому предпочитаем его
    return FONT_NAME


# ============================================================================
# ФУНКЦИИ ГЕНЕРАЦИИ WORD-ДОКУМЕНТОВ
# ============================================================================

def group_rows_by_construction(data_rows: List[List[str]], row_indices: List[int], 
                                construction_col: Optional[int]) -> List[Tuple[Optional[str], List[int]]]:
    """
    Группирует строки по конструкциям.
    
    Args:
        data_rows: Все строки данных из Excel
        row_indices: Индексы выбранных строк (0-based)
        construction_col: Индекс столбца с наименованием конструкций или None
        
    Returns:
        List[Tuple[Optional[str], List[int]]]: Список (название_конструкции, [индексы_строк])
    """
    if construction_col is None:
        # Если нет столбца конструкций, все строки в одну группу
        return [(None, row_indices)]
    
    # Группируем по конструкциям
    groups = {}
    for row_idx in row_indices:
        if row_idx < len(data_rows):
            row = data_rows[row_idx]
            if construction_col < len(row):
                construction_name = str(row[construction_col]).strip()
                if not construction_name:
                    construction_name = None
            else:
                construction_name = None
        else:
            construction_name = None
        
        if construction_name not in groups:
            groups[construction_name] = []
        groups[construction_name].append(row_idx)
    
    # Возвращаем в виде списка кортежей
    result = [(name, sorted(indices)) for name, indices in groups.items()]
    # Сортируем так, чтобы None был в конце
    result.sort(key=lambda x: (x[0] is None, x[0] or ""))
    return result


def create_word_document(data_rows: List[List[str]], row_indices: List[int], 
                        output_filename: str, source_file: str, column_mapping: dict,
                        construction_col: Optional[int]) -> Optional[Path]:
    """
    Создаёт Word-документ с таблицей из выбранных строк Excel.
    
    Args:
        data_rows: Все строки данных из Excel
        row_indices: Индексы выбранных строк (0-based)
        output_filename: Имя выходного файла
        source_file: Имя исходного Excel-файла (для информации)
        column_mapping: Маппинг столбцов {word_column_index: excel_column_index}
        construction_col: Индекс столбца с наименованием конструкций или None
        
    Returns:
        Path: Путь к созданному файлу или None при ошибке
    """
    try:
        # Создаём новый документ
        doc = Document()
        
        # Настройка страницы: A4, альбомная ориентация
        section = doc.sections[0]
        
        # Устанавливаем альбомную ориентацию
        section.orientation = WD_ORIENT.LANDSCAPE
        
        # Размеры A4 в альбомной ориентации: ширина (горизонталь) 11.69", высота (вертикаль) 8.27"
        # В альбомной ориентации ширина больше высоты
        section.page_width = Inches(11.69)   # Широкая сторона A4 (горизонталь)
        section.page_height = Inches(8.27)   # Короткая сторона A4 (вертикаль)
        
        # Устанавливаем поля
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        
        # Добавляем заголовок документа (опционально)
        title = doc.add_paragraph("Список дефектов и повреждений")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_format = title.runs[0].font
        title_format.name = get_font_name()
        title_format.size = Pt(14)
        title_format.bold = True
        
        # Добавляем информацию об источнике
        info = doc.add_paragraph(f"Источник: {source_file} | Строк: {len(row_indices)}")
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_format = info.runs[0].font
        info_format.name = get_font_name()
        info_format.size = Pt(9)
        info_format.italic = True
        
        # Добавляем пустую строку
        doc.add_paragraph()
        
        # Группируем строки по конструкциям
        construction_groups = group_rows_by_construction(data_rows, row_indices, construction_col)
        
        # Подсчитываем общее количество строк: 2 заголовка + строки данных + строки-заголовки конструкций
        num_data_rows = len(row_indices)
        num_construction_headers = len([g for g in construction_groups if g[0] is not None])
        num_cols = len(TABLE_HEADERS)
        num_rows_table = 2 + num_data_rows + num_construction_headers  # 2 заголовка + данные + заголовки конструкций
        
        table = doc.add_table(rows=num_rows_table, cols=num_cols)
        table.style = 'Table Grid'  # Базовый стиль с границами
        
        # Настройка ширины столбцов - первый столбец минимальной ширины по содержимому
        # № п/п - минимальная ширина (0.25"), остальные распределяются пропорционально
        column_widths = [0.25, 2.5, 1.5, 1.0, 2.5, 1.5]  # № п/п (мин), Наименование, Место, Фото, Параметр, Примечание
        total_width = sum(column_widths)
        
        for idx, col in enumerate(table.columns):
            if idx == 0:
                # Первый столбец "№ п/п" - минимальная ширина по содержимому
                # Устанавливаем минимальную ширину, которая будет автоматически подстраиваться
                col.width = Inches(0.25)
            else:
                # Остальные столбцы распределяются пропорционально
                col.width = Inches(column_widths[idx] / (total_width - column_widths[0]) * (9.27 - 0.25))
        
        # Заполняем первую строку заголовков таблицы (наименования столбцов)
        header_row1 = table.rows[0]
        for col_idx, header_text in enumerate(TABLE_HEADERS):
            cell = header_row1.cells[col_idx]
            cell.text = header_text
            
            # Форматирование заголовка
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = get_font_name()
                    run.font.size = FONT_SIZE
                    run.font.bold = True
            
            # Настройка ячейки (выравнивание по центру по вертикали)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        # Заполняем вторую строку заголовков (нумерация столбцов)
        header_row2 = table.rows[1]
        for col_idx, column_number in enumerate(COLUMN_NUMBERS):
            cell = header_row2.cells[col_idx]
            cell.text = column_number
            
            # Форматирование номера столбца
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = get_font_name()
                    run.font.size = FONT_SIZE
                    run.font.bold = True
            
            # Настройка ячейки (выравнивание по центру по вертикали)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        # Заполняем данные из выбранных строк с группировкой по конструкциям
        current_table_row = 2  # Начинаем с 2, так как 0 и 1 - это заголовки
        row_counter = 1  # Счётчик для нумерации строк в столбце "№ п/п"
        
        for construction_name, group_row_indices in construction_groups:
            # Если есть название конструкции, вставляем строку-заголовок
            if construction_name is not None:
                header_row = table.rows[current_table_row]
                
                # Объединяем ячейки через всю таблицу для заголовка конструкции
                # Объединяем с конца, чтобы избежать проблем с изменяющимся списком ячеек
                first_cell = header_row.cells[0]
                # Объединяем все остальные ячейки с первой
                for col_idx in range(num_cols - 1, 0, -1):
                    first_cell.merge(header_row.cells[col_idx])
                
                # Устанавливаем текст заголовка конструкции
                first_cell.text = construction_name
                
                # Форматирование строки-заголовка конструкции
                first_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for paragraph in first_cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = get_font_name()
                        run.font.size = Pt(12)  # Чуть крупнее для заголовка
                        run.font.bold = True
                first_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                
                current_table_row += 1
            
            # Заполняем строки данных для этой конструкции
            for data_row_idx in group_row_indices:
                row = table.rows[current_table_row]
                data_row = data_rows[data_row_idx]
                
                # Столбец 0: № п/п - порядковая нумерация строк
                row.cells[0].text = str(row_counter)
                
                # Столбец 1: Наименование дефекта/повреждения
                excel_col_idx = column_mapping.get(1)
                if excel_col_idx is not None and excel_col_idx < len(data_row):
                    row.cells[1].text = str(data_row[excel_col_idx])
                else:
                    row.cells[1].text = ""
                
                # Столбец 2: Место расположения - всегда пустой
                row.cells[2].text = ""
                
                # Столбец 3: Номер фото из фотоматериалов
                excel_col_idx = column_mapping.get(3)
                if excel_col_idx is not None and excel_col_idx < len(data_row):
                    row.cells[3].text = str(data_row[excel_col_idx])
                else:
                    row.cells[3].text = ""
                
                # Столбец 4: Контролируемый параметр дефекта
                excel_col_idx = column_mapping.get(4)
                if excel_col_idx is not None and excel_col_idx < len(data_row):
                    row.cells[4].text = str(data_row[excel_col_idx])
                else:
                    row.cells[4].text = ""
                
                # Столбец 5: Примечание
                excel_col_idx = column_mapping.get(5)
                if excel_col_idx is not None and excel_col_idx < len(data_row):
                    row.cells[5].text = str(data_row[excel_col_idx])
                else:
                    row.cells[5].text = ""
                
                # Форматирование всех ячеек строки
                for word_col_idx in range(num_cols):
                    cell = row.cells[word_col_idx]
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = get_font_name()
                            run.font.size = FONT_SIZE
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                
                row_counter += 1  # Увеличиваем счётчик порядковой нумерации
                current_table_row += 1
        
        # Настройка границ таблицы
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        tbl = table._tbl
        tblBorders = OxmlElement('w:tblBorders')
        
        borders = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
        for border_name in borders:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # 0.5pt = 4 единиц (1/8 pt)
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')  # Чёрный
            tblBorders.append(border)
        
        tbl.tblPr.append(tblBorders)
        
        # Сохраняем документ
        output_path = OUTPUT_DIR / output_filename
        
        # Проверяем, существует ли файл и можем ли мы его перезаписать
        if output_path.exists():
            try:
                # Пробуем удалить существующий файл, если он не открыт
                output_path.unlink()
            except PermissionError:
                # Файл открыт в другой программе, создаем файл с другим именем
                print(f"⚠ Файл {output_filename} уже открыт в другой программе.")
                print("   Создаю файл с другим именем...")
                
                # Генерируем уникальное имя файла с временной меткой
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                name_parts = output_filename.rsplit('.', 1)
                if len(name_parts) == 2:
                    new_filename = f"{name_parts[0]}_{timestamp}.{name_parts[1]}"
                else:
                    new_filename = f"{output_filename}_{timestamp}"
                
                output_path = OUTPUT_DIR / new_filename
        
        # Пытаемся сохранить документ
        try:
            doc.save(str(output_path))
            print(f"✓ Документ сохранён: {output_path}")
            return output_path
        except PermissionError as e:
            print(f"\n❌ ОШИБКА: Не удалось сохранить документ!")
            print(f"   Путь: {output_path}")
            print(f"\n💡 ВОЗМОЖНЫЕ ПРИЧИНЫ:")
            print("   1. Файл открыт в программе Word (закройте его)")
            print("   2. Нет прав на запись в папку output")
            print("   3. Файл используется другой программой")
            print(f"\n   Попробуйте:")
            print(f"   • Закрыть файл {output_filename} в Word, если он открыт")
            print(f"   • Закрыть все программы, использующие файл")
            print(f"   • Проверить права доступа к папке output")
            return None
        
    except Exception as e:
        print(f"❌ Ошибка при создании Word-документа: {e}")
        import traceback
        traceback.print_exc()
        return None


def open_document(file_path: Path) -> None:
    """
    Автоматически открывает созданный документ в системном приложении.
    
    Args:
        file_path: Путь к файлу для открытия
    """
    try:
        if sys.platform == "win32":
            # Windows
            os.startfile(str(file_path))
        elif sys.platform == "darwin":
            # macOS
            subprocess.call(["open", str(file_path)])
        else:
            # Linux или другие ОС
            subprocess.call(["xdg-open", str(file_path)])
        
        print(f"✓ Документ открыт: {file_path.name}")
        
    except Exception as e:
        print(f"⚠ Не удалось автоматически открыть документ: {e}")
        print(f"   Откройте его вручную: {file_path}")


# ============================================================================
# ГЛАВНАЯ ФУНКЦИЯ
# ============================================================================

def main():
    """Основная функция скрипта."""
    print("="*60)
    print("ГЕНЕРАТОР WORD-ДОКУМЕНТОВ ИЗ EXCEL")
    print("="*60)
    print(f"Рабочая директория: {BASE_DIR}")
    
    # Создаём необходимые директории
    create_directories()
    
    # Получаем список Excel-файлов
    excel_files = get_excel_files()
    
    # Выбираем файл
    selected_file = select_excel_file(excel_files)
    if selected_file is None:
        print("\n❌ Файл не выбран. Завершение работы.")
        return
    
    # Читаем данные из Excel (возвращает данные, имя листа, маппинг столбцов и индекс столбца конструкций)
    data, sheet_name, column_mapping, construction_col = read_excel_file(selected_file)
    if not data:
        print("\n❌ Не удалось загрузить данные из Excel-файла.")
        return
    
    # Сначала выбираем конструкции для обработки (если столбец конструкций найден)
    selected_constructions = select_constructions(data, construction_col)
    
    # Предварительная фильтрация данных по выбранным конструкциям (если выбраны)
    filtered_data = data
    if selected_constructions is not None and construction_col is not None:
        # Фильтруем данные по конструкциям перед выбором строк
        filtered_data = []
        for row in data:
            if construction_col < len(row):
                construction_name = str(row[construction_col]).strip()
                if construction_name in selected_constructions:
                    filtered_data.append(row)
        print(f"\n✓ Данные предварительно отфильтрованы по выбранным конструкциям.")
        print(f"   Доступно {len(filtered_data)} строк с дефектами для выбранных конструкций.")
        if not filtered_data:
            print("\n❌ После фильтрации по конструкциям не осталось строк. Завершение работы.")
            return
    else:
        filtered_data = data
    
    # Выбираем строки для обработки из отфильтрованных данных
    selected_row_indices = select_rows(filtered_data)
    if not selected_row_indices:
        print("\n❌ Не выбрано ни одной строки. Завершение работы.")
        return
    
    # Если использовали отфильтрованные данные, нужно пересчитать индексы
    if selected_constructions is not None and construction_col is not None:
        # Находим исходные индексы строк в оригинальных данных
        original_indices = []
        filtered_idx = 0
        for orig_idx, row in enumerate(data):
            if construction_col < len(row):
                construction_name = str(row[construction_col]).strip()
                if construction_name in selected_constructions:
                    if filtered_idx in selected_row_indices:
                        original_indices.append(orig_idx)
                    filtered_idx += 1
        selected_row_indices = original_indices
    
    # Генерируем имя файла
    if len(selected_row_indices) == 1:
        row_label = str(selected_row_indices[0] + 1)
    else:
        first_row = selected_row_indices[0] + 1
        last_row = selected_row_indices[-1] + 1
        row_label = f"{first_row}-{last_row}"
    
    output_filename = f"Дефекты_выборка_{row_label}.docx"
    
    # Создаём Word-документ (передаём маппинг столбцов и индекс столбца конструкций)
    output_path = create_word_document(
        data, 
        selected_row_indices, 
        output_filename,
        selected_file.name,
        column_mapping,
        construction_col
    )
    
    if output_path:
        # Открываем документ
        open_document(output_path)
        print("\n" + "="*60)
        print("✓ ГОТОВО! Документ создан и открыт.")
        print("="*60)
    else:
        print("\n❌ Не удалось создать документ.")


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n\n⚠ Операция прервана пользователем.")
        sys.exit(0)
    except Exception as e:
        print(f"\n❌ Критическая ошибка: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

